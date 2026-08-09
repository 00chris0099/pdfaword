import logging
from pathlib import Path

logger = logging.getLogger("pdfforge.quality")


def check_conversion_quality(original_pdf: str, output_docx: str) -> dict:
    try:
        import pymupdf
        from docx import Document

        # Analyze original
        pdf_doc = pymupdf.open(original_pdf)
        pdf_text_length = 0
        pdf_images = 0
        pdf_pages = len(pdf_doc)

        for page in pdf_doc:
            pdf_text_length += len(page.get_text())
            pdf_images += len(page.get_images())
        pdf_doc.close()

        # Analyze output
        docx_doc = Document(output_docx)
        docx_text_length = 0
        docx_paragraphs = len(docx_doc.paragraphs)
        docx_tables = len(docx_doc.tables)
        docx_images = 0

        for rel in docx_doc.part.rels.values():
            if "image" in rel.reltype:
                docx_images += 1

        for para in docx_doc.paragraphs:
            docx_text_length += len(para.text)

        # Calculate text preservation ratio
        text_ratio = docx_text_length / max(pdf_text_length, 1)

        quality_score = min(1.0, text_ratio)
        issues = []

        if text_ratio < 0.8:
            issues.append(f"Texto preservado: {text_ratio:.0%} (esperado: >80%)")

        if pdf_images > 0 and docx_images == 0:
            issues.append("Imágenes no preservadas")

        return {
            "quality_score": round(quality_score, 2),
            "text_preservation_ratio": round(text_ratio, 2),
            "pdf_pages": pdf_pages,
            "pdf_text_length": pdf_text_length,
            "pdf_images": pdf_images,
            "docx_paragraphs": docx_paragraphs,
            "docx_tables": docx_tables,
            "docx_images": docx_images,
            "issues": issues,
            "passed": quality_score >= 0.8 and len(issues) == 0,
        }

    except Exception as e:
        logger.exception(f"Quality check failed: {e}")
        return {"quality_score": 0, "issues": [str(e)], "passed": False}
