import io
import time
import uuid
import shutil
import zipfile
from datetime import datetime
from pathlib import Path
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, BackgroundTasks, status
from fastapi.responses import StreamingResponse, FileResponse
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from app.database import get_db
from app.models import User, Conversion, ConversionStatus, BatchJob, BatchItem
from app.schemas import (
    ConversionResponse,
    ConversionUploadResponse,
    BatchUploadResponse,
    BatchStatusResponse,
    BatchItemResponse,
    BatchDetailResponse,
    ErrorResponse,
    HTMLConvertResponse,
)
from app.auth import get_current_user
from app.config import settings
from app.services.converter import run_conversion
from app.services.batch import process_batch

router = APIRouter(prefix="/api/convert", tags=["convert"])

ALLOWED_EXTENSIONS = {".pdf"}
MAX_FILE_SIZE = settings.MAX_FILE_SIZE_MB * 1024 * 1024


def validate_pdf(file: UploadFile) -> None:
    if not file.filename:
        raise HTTPException(status_code=400, detail="Nombre de archivo requerido")
    ext = Path(file.filename).suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=400, detail="Solo se permiten archivos PDF")


async def save_upload(file: UploadFile, dest: Path) -> int:
    total_size = 0
    with open(dest, "wb") as buffer:
        while chunk := await file.read(8192):
            total_size += len(chunk)
            if total_size > MAX_FILE_SIZE:
                buffer.close()
                dest.unlink(missing_ok=True)
                raise HTTPException(
                    status_code=413,
                    detail=f"El archivo excede el límite de {settings.MAX_FILE_SIZE_MB}MB",
                )
            buffer.write(chunk)
    return total_size


@router.post("", response_model=ConversionUploadResponse, status_code=status.HTTP_202_ACCEPTED)
async def upload_and_convert(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    translate_to: str | None = None,
    force_ocr: bool = False,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    # Validate
    validate_pdf(file)

    # Check credits (-1 = unlimited)
    if user.credits_remaining == 0:
        raise HTTPException(status_code=402, detail="Sin créditos disponibles. Actualiza tu plan.")

    # Create conversion record
    conversion = Conversion(
        user_id=user.id,
        original_filename=file.filename,
        original_size=0,
        status=ConversionStatus.PENDING,
        file_path_pdf="",
        translation_lang=translate_to,
    )
    db.add(conversion)
    await db.commit()
    await db.refresh(conversion)

    # Save PDF to disk
    file_id = f"{conversion.id}_{uuid.uuid4().hex[:8]}"
    pdf_path = settings.UPLOAD_DIR / f"{file_id}.pdf"
    file_size = await save_upload(file, pdf_path)

    # Update record with file info
    conversion.original_size = file_size
    conversion.file_path_pdf = str(pdf_path)
    await db.commit()

    # Deduct credit
    user.credits_remaining -= 1
    await db.commit()

    # Launch background conversion
    background_tasks.add_task(run_conversion, conversion.id, force_ocr)

    return ConversionUploadResponse(
        id=conversion.id,
        message="PDF recibido. Conversión en progreso.",
        status=conversion.status,
        original_filename=conversion.original_filename,
    )


@router.get("/{conversion_id}/status", response_model=ConversionResponse)
async def get_conversion_status(
    conversion_id: int,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(Conversion).where(
            Conversion.id == conversion_id,
            Conversion.user_id == user.id,
        )
    )
    conversion = result.scalar_one_or_none()
    if not conversion:
        raise HTTPException(status_code=404, detail="Conversión no encontrada")
    return ConversionResponse.model_validate(conversion)


@router.get("/{conversion_id}/download")
async def download_conversion(
    conversion_id: int,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(Conversion).where(
            Conversion.id == conversion_id,
            Conversion.user_id == user.id,
        )
    )
    conversion = result.scalar_one_or_none()
    if not conversion:
        raise HTTPException(status_code=404, detail="Conversión no encontrada")

    if conversion.status != ConversionStatus.COMPLETED:
        raise HTTPException(status_code=400, detail="La conversión aún no está completa")

    if not conversion.file_path_docx or not Path(conversion.file_path_docx).exists():
        raise HTTPException(status_code=404, detail="Archivo DOCX no encontrado")

    from fastapi.responses import FileResponse
    docx_name = Path(conversion.original_filename).stem + ".docx"
    return FileResponse(
        path=conversion.file_path_docx,
        filename=docx_name,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@router.delete("/{conversion_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_conversion(
    conversion_id: int,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(Conversion).where(
            Conversion.id == conversion_id,
            Conversion.user_id == user.id,
        )
    )
    conversion = result.scalar_one_or_none()
    if not conversion:
        raise HTTPException(status_code=404, detail="Conversión no encontrada")

    # Delete files
    if conversion.file_path_pdf:
        Path(conversion.file_path_pdf).unlink(missing_ok=True)
    if conversion.file_path_docx:
        Path(conversion.file_path_docx).unlink(missing_ok=True)

    await db.delete(conversion)
    await db.commit()


# ── Batch Upload ──────────────────────────────────────

@router.post("/batch", response_model=BatchUploadResponse, status_code=status.HTTP_202_ACCEPTED)
async def upload_batch(
    background_tasks: BackgroundTasks,
    files: list[UploadFile] = File(..., max_length=20),
    translate_to: str | None = None,
    force_ocr: bool = False,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    if len(files) < 2:
        raise HTTPException(status_code=400, detail="Batch requiere al menos 2 archivos")
    if len(files) > 20:
        raise HTTPException(status_code=400, detail="Máximo 20 archivos por batch")

    if user.credits_remaining != -1 and user.credits_remaining < len(files):
        raise HTTPException(
            status_code=402,
            detail=f"Créditos insuficientes. Necesitas {len(files)}, tienes {user.credits_remaining}",
        )

    # Create batch
    batch = BatchJob(
        user_id=user.id,
        status=ConversionStatus.PENDING,
        total_files=len(files),
        translation_lang=translate_to,
    )
    db.add(batch)
    await db.commit()
    await db.refresh(batch)

    saved_files = []
    for f in files:
        validate_pdf(f)
        file_id = f"{batch.id}_{uuid.uuid4().hex[:8]}"
        pdf_path = settings.UPLOAD_DIR / f"{file_id}.pdf"
        file_size = await save_upload(f, pdf_path)
        saved_files.append((pdf_path, f.filename, file_size))

    # Create conversions + batch items
    for pdf_path, filename, file_size in saved_files:
        conversion = Conversion(
            user_id=user.id,
            original_filename=filename,
            original_size=file_size,
            status=ConversionStatus.PENDING,
            file_path_pdf=str(pdf_path),
            translation_lang=translate_to,
        )
        db.add(conversion)
        await db.commit()
        await db.refresh(conversion)

        item = BatchItem(
            batch_id=batch.id,
            conversion_id=conversion.id,
            filename=filename,
            status=ConversionStatus.PENDING,
        )
        db.add(item)

    await db.commit()

    # Deduct credits (enterprise = unlimited)
    if user.credits_remaining != -1:
        user.credits_remaining -= len(files)
        await db.commit()

    # Launch background processing
    background_tasks.add_task(process_batch, batch.id)

    return BatchUploadResponse(
        id=batch.id,
        message=f"Batch de {len(files)} archivos recibido. Procesando en paralelo.",
        total_files=len(files),
        status=batch.status,
    )


@router.get("/batch/{batch_id}", response_model=BatchDetailResponse)
async def get_batch_status(
    batch_id: int,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(BatchJob).where(BatchJob.id == batch_id, BatchJob.user_id == user.id)
    )
    batch = result.scalar_one_or_none()
    if not batch:
        raise HTTPException(status_code=404, detail="Batch no encontrado")

    items_result = await db.execute(
        select(BatchItem).where(BatchItem.batch_id == batch_id)
    )
    items = items_result.scalars().all()

    item_responses = []
    for item in items:
        conv = None
        if item.conversion_id:
            conv_result = await db.execute(
                select(Conversion).where(Conversion.id == item.conversion_id)
            )
            conv = conv_result.scalar_one_or_none()
        item_responses.append(
            BatchItemResponse(
                id=item.id,
                filename=item.filename,
                status=item.status,
                conversion_id=item.conversion_id,
                output_size=conv.output_size if conv else None,
            )
        )

    return BatchDetailResponse(
        id=batch.id,
        status=batch.status,
        total_files=batch.total_files,
        completed_files=batch.completed_files,
        failed_files=batch.failed_files,
        created_at=batch.created_at,
        completed_at=batch.completed_at,
        items=item_responses,
    )


@router.get("/batch/{batch_id}/download")
async def download_batch_zip(
    batch_id: int,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(BatchJob).where(BatchJob.id == batch_id, BatchJob.user_id == user.id)
    )
    batch = result.scalar_one_or_none()
    if not batch:
        raise HTTPException(status_code=404, detail="Batch no encontrado")

    if batch.status != ConversionStatus.COMPLETED:
        raise HTTPException(status_code=400, detail="El batch aún no está completo")

    items_result = await db.execute(
        select(BatchItem).where(BatchItem.batch_id == batch_id)
    )
    items = items_result.scalars().all()

    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        for item in items:
            if item.conversion_id:
                conv_result = await db.execute(
                    select(Conversion).where(Conversion.id == item.conversion_id)
                )
                conv = conv_result.scalar_one_or_none()
                if conv and conv.file_path_docx and Path(conv.file_path_docx).exists():
                    docx_name = Path(conv.original_filename).stem + ".docx"
                    zf.write(conv.file_path_docx, docx_name)

    zip_buffer.seek(0)
    return StreamingResponse(
        zip_buffer,
        media_type="application/zip",
        headers={"Content-Disposition": f"attachment; filename=pdfforge_batch_{batch_id}.zip"},
    )


# ── Credit Status ─────────────────────────────────────

@router.get("/credits")
async def get_credit_status(user: User = Depends(get_current_user)):
    return {
        "credits_remaining": user.credits_remaining,
        "plan": user.plan.value,
        "unlimited": user.credits_remaining == -1,
    }


# ── HTML to DOCX ────────────────────────────────────

@router.post("/html", response_model=HTMLConvertResponse, status_code=status.HTTP_202_ACCEPTED)
async def convert_html_to_docx(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    html_content = (await file.read()).decode("utf-8")
    safe_filename = Path(file.filename).stem[:50] if file.filename else "documento"
    safe_filename = "".join(c for c in safe_filename if c.isalnum() or c in "-_ ").strip() or "documento"

    conversion = Conversion(
        user_id=user.id,
        original_filename=f"{safe_filename}.html",
        original_size=len(html_content.encode("utf-8")),
        status=ConversionStatus.PENDING,
        file_path_pdf="",
    )
    db.add(conversion)
    await db.commit()
    await db.refresh(conversion)

    html_path = settings.UPLOAD_DIR / f"{conversion.id}_{uuid.uuid4().hex[:8]}.html"
    html_path.write_text(html_content, encoding="utf-8")

    output_filename = f"{safe_filename}.docx"
    output_path = settings.OUTPUT_DIR / output_filename

    background_tasks.add_task(_process_html_to_docx, conversion.id, str(html_path), str(output_path), safe_filename)

    return HTMLConvertResponse(
        id=conversion.id,
        message="HTML recibido. Convirtiendo a Word...",
        status=conversion.status,
        original_filename=conversion.original_filename,
    )


def _process_html_to_docx(conversion_id: int, html_path: str, docx_path: str, filename: str):
    import time

    start = time.time()
    try:
        with open(html_path, "r", encoding="utf-8") as f:
            html_content = f.read()

        import mammoth

        result = mammoth.convert_to_html(html_content)
        html_output = result.value
        messages = result.messages

        if not html_output or len(html_output.strip()) < 10:
            raise ValueError("Mammoth no pudo convertir el HTML. Usando fallback.")

        html_bytes = html_output.encode("utf-8")
        doc_result = mammoth.convert_to_docx(html_bytes)
        docx_bytes = doc_result.value

        with open(docx_path, "wb") as f:
            f.write(docx_bytes)

        elapsed = time.time() - start
        import asyncio
        asyncio.run(_update_conversion_done(conversion_id, docx_path, elapsed))

    except Exception as e:
        try:
            _fallback_html_to_docx(conversion_id, html_path, docx_path, start)
        except Exception as fallback_error:
            import asyncio
            asyncio.run(_update_conversion_failed(conversion_id, f"Error: {e}. Fallback: {fallback_error}"))


def _fallback_html_to_docx(conversion_id: int, html_path: str, docx_path: str, start_time: float):
    import re
    from docx import Document
    from docx.shared import Pt

    with open(html_path, "r", encoding="utf-8") as f:
        html_content = f.read()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    html_content = re.sub(r"<style[^>]*>.*?</style>", "", html_content, flags=re.DOTALL | re.IGNORECASE)
    html_content = re.sub(r"<script[^>]*>.*?</script>", "", html_content, flags=re.DOTALL | re.IGNORECASE)

    body_match = re.search(r"<body[^>]*>(.*?)</body>", html_content, re.DOTALL | re.IGNORECASE)
    if body_match:
        html_content = body_match.group(1)

    table_pattern = re.compile(r"<table[^>]*>(.*?)</table>", re.DOTALL | re.IGNORECASE)
    tables = table_pattern.findall(html_content)
    for table_html in tables:
        _add_html_table(doc, "<table>" + table_html + "</table>")
        html_content = html_content.replace("<table>" + table_html + "</table>", "")

    lines = re.split(r"\n|<br\s*/?>|</p>|</div>|</h[1-6]>", html_content)
    for line in lines:
        line = line.strip()
        if not line:
            continue
        clean = re.sub(r"<[^>]+>", "", line).strip()
        if not clean:
            continue
        if re.match(r"<h[1-6]", line, re.IGNORECASE):
            level = int(re.search(r"h(\d)", line, re.IGNORECASE).group(1))
            doc.add_heading(clean, level=min(level, 4))
        elif re.match(r"<li", line, re.IGNORECASE):
            doc.add_paragraph(clean, style="List Bullet")
        else:
            doc.add_paragraph(clean)

    doc.save(docx_path)

    elapsed = time.time() - start_time
    import asyncio
    asyncio.run(_update_conversion_done(conversion_id, docx_path, elapsed))


async def _update_conversion_done(conversion_id: int, docx_path: str, elapsed: float):
    from app.database import async_session
    from sqlalchemy import select
    from app.models import Conversion, ConversionStatus
    from pathlib import Path
    from datetime import datetime

    async with async_session() as db:
        result = await db.execute(select(Conversion).where(Conversion.id == conversion_id))
        conv = result.scalar_one_or_none()
        if conv:
            conv.status = ConversionStatus.COMPLETED
            conv.status_message = "Conversión completada"
            conv.output_size = Path(docx_path).stat().st_size if Path(docx_path).exists() else 0
            conv.file_path_docx = docx_path
            conv.conversion_time_seconds = round(elapsed, 2)
            conv.completed_at = datetime.utcnow()
            await db.commit()


async def _update_conversion_failed(conversion_id: int, error: str):
    from app.database import async_session
    from sqlalchemy import select
    from app.models import Conversion, ConversionStatus

    async with async_session() as db:
        result = await db.execute(select(Conversion).where(Conversion.id == conversion_id))
        conv = result.scalar_one_or_none()
        if conv:
            conv.status = ConversionStatus.FAILED
            conv.status_message = error
            await db.commit()


def _strip_tags(html: str) -> str:
    import re
    clean = re.sub(r"<[^>]+>", "", html)
    from html import unescape
    return unescape(clean).strip()


def _add_html_table(doc, html_table: str) -> None:
    import re

    rows = re.findall(r"<tr[^>]*>(.*?)</tr>", html_table, re.DOTALL)
    if not rows:
        return

    cols = 0
    parsed_rows = []
    for row_html in rows:
        cells = re.findall(r"<t[dh][^>]*>(.*?)</t[dh]>", row_html, re.DOTALL)
        parsed_rows.append([_strip_tags(c) for c in cells])
        cols = max(cols, len(cells))

    if cols == 0:
        return

    table = doc.add_table(rows=len(parsed_rows), cols=cols, style="Table Grid")
    for i, row_data in enumerate(parsed_rows):
        for j, cell_text in enumerate(row_data):
            if j < cols:
                table.cell(i, j).text = cell_text

    doc.add_paragraph()
