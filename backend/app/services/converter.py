import time
import logging
from pathlib import Path
from datetime import datetime
from sqlalchemy import select
from app.database import async_session
from app.models import Conversion, ConversionStatus
from app.config import settings
from app.utils.pdf_analyzer import analyze_pdf
from app.services.ocr_engine import ocr_engine
from app.services.docx_builder import convert_pdf_to_docx, build_docx_from_blocks
from app.services.translator import translate_docx

logger = logging.getLogger("pdfforge.converter")


def _append_headers_footers_to_docx(docx_path: str, pdf_path: str) -> None:
    """Extract headers and footers from PDF and append to DOCX if missing."""
    try:
        import pymupdf
        from docx import Document
        from docx.shared import Pt, RGBColor

        pdf_doc = pymupdf.open(pdf_path)
        doc = Document(docx_path)

        existing_text = "\n".join(p.text for p in doc.paragraphs)

        for page_num in range(len(pdf_doc)):
            page = pdf_doc[page_num]
            page_height = page.rect.height
            page_width = page.rect.width

            # Header: top 10% of page
            header_area = pymupdf.Rect(0, 0, page_width, page_height * 0.10)
            header_text = page.get_text("text", clip=header_area).strip()

            # Footer: bottom 12% of page
            footer_area = pymupdf.Rect(0, page_height * 0.88, page_width, page_height)
            footer_text = page.get_text("text", clip=footer_area).strip()

            # Add header if not in existing text
            if header_text:
                for line in header_text.split("\n"):
                    line = line.strip()
                    if line and line not in existing_text:
                        p = doc.add_paragraph()
                        p.alignment = 1  # center
                        run = p.add_run(line)
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(100, 100, 100)
                        existing_text += "\n" + line

            # Add footer if not in existing text
            if footer_text:
                for line in footer_text.split("\n"):
                    line = line.strip()
                    if line and line not in existing_text:
                        p = doc.add_paragraph()
                        p.alignment = 1  # center
                        run = p.add_run(line)
                        run.font.size = Pt(8)
                        run.font.color.rgb = RGBColor(128, 128, 128)
                        existing_text += "\n" + line

        doc.save(docx_path)
        pdf_doc.close()
        logger.info(f"Appended headers/footers to {docx_path}")

    except Exception as e:
        logger.warning(f"Header/footer extraction failed: {e}")


async def run_conversion(conversion_id: int, force_ocr: bool = False):
    start_time = time.time()
    async with async_session() as db:
        result = await db.execute(select(Conversion).where(Conversion.id == conversion_id))
        conversion = result.scalar_one_or_none()
        if not conversion:
            logger.error(f"Conversion {conversion_id} not found")
            return

        try:
            # Step 1: Analyzing
            conversion.status = ConversionStatus.ANALYZING
            conversion.status_message = "Analizando PDF..."
            await db.commit()

            pdf_path = Path(conversion.file_path_pdf)
            analysis = analyze_pdf(str(pdf_path))
            conversion.pages_count = analysis["pages"]
            conversion.has_images = analysis["has_images"]
            conversion.has_tables = analysis["has_tables"]
            conversion.is_scanned = analysis["is_scanned"]
            await db.commit()

            logger.info(f"Conversion {conversion_id}: analyzed - {analysis}")

            # Step 2: Convert
            output_filename = pdf_path.stem + ".docx"
            output_path = settings.OUTPUT_DIR / output_filename

            use_ocr = force_ocr or analysis["is_scanned"]

            if use_ocr:
                conversion.status = ConversionStatus.OCR_PROCESSING
                conversion.status_message = "Extrayendo texto con OCR..."
                conversion.ocr_used = True
                await db.commit()

                logger.info(f"Conversion {conversion_id}: OCR extraction")
                blocks = ocr_engine.process(str(pdf_path))
                conversion.status_message = f"Construyendo Word ({len(blocks)} bloques)..."
                await db.commit()
                build_docx_from_blocks(blocks, str(output_path), pdf_path=str(pdf_path))
            else:
                conversion.status = ConversionStatus.CONVERTING
                conversion.status_message = "Convirtiendo PDF a Word (preservando formato)..."
                conversion.ocr_used = False
                await db.commit()

                logger.info(f"Conversion {conversion_id}: pdf2docx conversion")
                convert_pdf_to_docx(str(pdf_path), str(output_path))

                # Post-process: append missing headers and footers
                _append_headers_footers_to_docx(str(output_path), str(pdf_path))

            # Step 3: Translation (optional)
            if conversion.translation_lang:
                conversion.status = ConversionStatus.TRANSLATING
                conversion.status_message = f"Traduciendo a {conversion.translation_lang}..."
                await db.commit()

                logger.info(f"Conversion {conversion_id}: translating to {conversion.translation_lang}")
                translate_docx(str(output_path), conversion.translation_lang)

            # Step 4: Done
            output_size = output_path.stat().st_size if output_path.exists() else 0
            elapsed = time.time() - start_time

            conversion.status = ConversionStatus.COMPLETED
            conversion.status_message = "Conversión completada"
            conversion.output_size = output_size
            conversion.file_path_docx = str(output_path)
            conversion.conversion_time_seconds = round(elapsed, 2)
            conversion.completed_at = datetime.utcnow()
            await db.commit()

            logger.info(f"Conversion {conversion_id}: completed in {elapsed:.2f}s")

        except Exception as e:
            elapsed = time.time() - start_time
            conversion.status = ConversionStatus.FAILED
            conversion.status_message = str(e)
            conversion.conversion_time_seconds = round(elapsed, 2)
            await db.commit()
            logger.exception(f"Conversion {conversion_id} failed: {e}")
