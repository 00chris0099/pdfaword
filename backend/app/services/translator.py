import logging
from pathlib import Path

logger = logging.getLogger("pdfforge.translator")


def translate_docx(docx_path: str, target_lang: str) -> None:
    try:
        from docx import Document
        from deep_translator import GoogleTranslator

        doc = Document(docx_path)
        translator = GoogleTranslator(source="auto", target=target_lang)

        # Translate paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                translated = translator.translate(paragraph.text)
                for run in paragraph.runs:
                    run.text = ""

                if paragraph.runs:
                    paragraph.runs[0].text = translated
                else:
                    paragraph.text = translated

        # Translate tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        translated = translator.translate(cell.text)
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.text = ""
                            if paragraph.runs:
                                paragraph.runs[0].text = translated
                            else:
                                paragraph.text = translated

        doc.save(docx_path)
        logger.info(f"Translated {docx_path} to {target_lang}")

    except ImportError:
        logger.error("deep-translator not installed")
        raise
    except Exception as e:
        logger.exception(f"Translation failed: {e}")
        raise
