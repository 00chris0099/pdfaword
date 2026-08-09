from __future__ import annotations
import logging
from pathlib import Path
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from docx import Document

from app.services.ocr_engine import RichBlock, RichRun

logger = logging.getLogger("pdfforge.docx_builder")

HEADING_SIZES = {1: 26, 2: 22, 3: 18, 4: 14}

HEADING_SPACING = {
    1: (24, 12),
    2: (20, 10),
    3: (16, 8),
    4: (12, 6),
}

LIST_BULLET_STYLE = "List Bullet"
LIST_NUMBER_STYLE = "List Bullet"


def convert_pdf_to_docx(pdf_path: str, docx_path: str) -> None:
    try:
        from pdf2docx import Converter
        cv = Converter(pdf_path)

        settings = {
            "ocr": 0,
            "ignore_page_error": True,
            "parse_lattice_table": True,
            "parse_stream_table": True,
            "clip_image_res_ratio": 4.0,
            "list_not_table": True,
            "connected_border_tolerance": 0.5,
            "min_border_clearance": 2.0,
            "float_image_ignorable_gap": 5.0,
            "line_break_width_ratio": 0.5,
            "new_paragraph_free_space_ratio": 0.85,
        }

        cv.convert(docx_path, **settings)
        cv.close()
        logger.info(f"pdf2docx converted {pdf_path} -> {docx_path}")
    except Exception as e:
        logger.exception(f"pdf2docx failed: {e}")
        raise


def build_docx_from_blocks(blocks: list[RichBlock], docx_path: str, pdf_path: str = "") -> None:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0, 0, 0)

    pf = style.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(4)
    pf.line_spacing = 1.15

    _set_page_size(doc, pdf_path)

    for block in blocks:
        try:
            if block.type == "heading":
                _add_heading(doc, block)
            elif block.type == "list_item":
                _add_list_item(doc, block)
            elif block.type == "table":
                _add_table(doc, block)
            elif block.type == "image":
                _add_image(doc, block)
            elif block.type == "paragraph":
                _add_paragraph(doc, block)
        except Exception as e:
            logger.warning(f"Failed to add block {block.type}: {e}")
            _add_fallback_paragraph(doc, block)

    doc.save(docx_path)
    logger.info(f"Built rich DOCX: {docx_path} ({len(blocks)} blocks)")


def _set_page_size(doc: "Document", pdf_path: str) -> None:
    if not pdf_path:
        return

    try:
        import pymupdf
        pdf_doc = pymupdf.open(pdf_path)
        if len(pdf_doc) > 0:
            page = pdf_doc[0]
            rect = page.rect
            width_inches = rect.width / 72
            height_inches = rect.height / 72

            section = doc.sections[0]
            section.page_width = Inches(width_inches)
            section.page_height = Inches(height_inches)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)
            section.top_margin = Inches(0.6)
            section.bottom_margin = Inches(0.6)
        pdf_doc.close()
    except Exception as e:
        logger.warning(f"Could not set page size: {e}")


def _add_heading(doc: "Document", block: RichBlock) -> None:
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    level = min(max(block.level, 1), 4)
    size = HEADING_SIZES.get(level, 14)

    p = doc.add_paragraph()
    p.style = doc.styles[f"Heading {level}"]
    p.alignment = _get_alignment(block.alignment)

    before, after = HEADING_SPACING.get(level, (12, 6))
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)

    if not block.runs:
        return

    for run_data in block.runs:
        run = p.add_run(run_data.text)
        run.font.size = Pt(size)
        run.font.bold = True
        run.font.name = run_data.font_name if run_data.font_name else "Calibri"
        if run_data.color and run_data.color != (0, 0, 0):
            run.font.color.rgb = RGBColor(*run_data.color)


def _add_paragraph(doc: "Document", block: RichBlock) -> None:
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    p = doc.add_paragraph()
    p.alignment = _get_alignment(block.alignment)

    if not block.runs:
        return

    prev_bold = None
    prev_italic = None
    prev_color = None
    current_text = ""

    def flush():
        nonlocal current_text, prev_bold, prev_italic, prev_color
        if current_text:
            run = p.add_run(current_text)
            run.font.name = "Calibri"
            if prev_bold is not None:
                run.font.bold = prev_bold
            if prev_italic is not None:
                run.font.italic = prev_italic
            if prev_color is not None and prev_color != (0, 0, 0):
                run.font.color.rgb = RGBColor(*prev_color)
            current_text = ""

    for run_data in block.runs:
        new_bold = run_data.bold
        new_italic = run_data.italic
        new_color = run_data.color

        if (
            new_bold != prev_bold
            or new_italic != prev_italic
            or new_color != prev_color
        ):
            flush()
            prev_bold = new_bold
            prev_italic = new_italic
            prev_color = new_color

        current_text += run_data.text

    flush()


def _add_list_item(doc: "Document", block: RichBlock) -> None:
    from docx.shared import Pt, RGBColor

    combined_text = "".join(r.text for r in block.runs)
    cleaned = _clean_list_prefix(combined_text)

    p = doc.add_paragraph(style=LIST_BULLET_STYLE)

    if block.indent_level > 0:
        p.paragraph_format.left_indent = Pt(18 * block.indent_level)

    if not block.runs:
        p.add_run(cleaned)
        return

    runs_to_add = []
    remaining = cleaned
    for run_data in block.runs:
        text = run_data.text
        if remaining.startswith(text):
            remaining = remaining[len(text):]
            runs_to_add.append((text, run_data))
        elif text.strip():
            runs_to_add.append((text, run_data))

    if not runs_to_add:
        p.add_run(cleaned)
        return

    for text, run_data in runs_to_add:
        run = p.add_run(text)
        run.font.name = run_data.font_name or "Calibri"
        run.font.size = Pt(run_data.font_size) if run_data.font_size else Pt(11)
        run.font.bold = run_data.bold
        run.font.italic = run_data.italic
        run.font.underline = run_data.underline
        if run_data.color and run_data.color != (0, 0, 0):
            run.font.color.rgb = RGBColor(*run_data.color)


def _clean_list_prefix(text: str) -> str:
    import re
    return re.sub(r"^\s*[\u2022\u2023\u25E6\u2043\u2219\-\*\+]+\s*", "", text)


def _add_table(doc: "Document", block: RichBlock) -> None:
    from docx.shared import Pt, RGBColor, Inches

    if not block.table_data or not block.table_data[0]:
        return

    rows = len(block.table_data)
    cols = block.table_col_count or len(block.table_data[0])
    cols = max(cols, 1)

    table = doc.add_table(rows=rows, cols=cols, style="Table Grid")

    for i, row_data in enumerate(block.table_data):
        for j, cell_text in enumerate(row_data):
            if j < cols:
                cell = table.cell(i, j)
                cell.text = str(cell_text) if cell_text else ""
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(2)
                    paragraph.paragraph_format.space_after = Pt(2)
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                        run.font.name = "Calibri"

    doc.add_paragraph()


def _add_image(doc: "Document", block: RichBlock) -> None:
    from docx.shared import Inches, Pt

    if not block.image_path or not Path(block.image_path).exists():
        return

    try:
        max_width = Inches(5.5)
        img_width = block.image_width
        img_height = block.image_height

        if img_width > 0 and img_height > 0:
            aspect = img_height / img_width
            width = min(max_width, Inches(6))
            height = Inches(width.inches * aspect)
            if height.inches > 8:
                height = Inches(8)
                width = Inches(height.inches / aspect)
        else:
            width = max_width
            height = Inches(3)

        p = doc.add_paragraph()
        p.alignment = 1  # center
        run = p.add_run()
        run.add_picture(block.image_path, width=width, height=height)

        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(4)
        spacer.paragraph_format.space_after = Pt(4)

    except Exception as e:
        logger.warning(f"Failed to add image: {e}")


def _add_fallback_paragraph(doc: "Document", block: RichBlock) -> None:
    combined = "".join(r.text for r in block.runs)
    if combined.strip():
        p = doc.add_paragraph()
        run = p.add_run(combined)
        run.font.size = Pt(11)
        run.font.name = "Calibri"


def _get_alignment(alignment: str):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    mapping = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    return mapping.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)
